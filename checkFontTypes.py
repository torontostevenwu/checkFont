import streamlit as st
from docx import Document

def get_fonts_from_docx(docx_file):
    doc = Document(docx_file)
    fonts = set()

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name:
                            fonts.add(run.font.name)

    return fonts

st.title("Font Types in Microsoft Word Document")

uploaded_file = st.file_uploader("Upload a Word document", type=["docx"])

if uploaded_file is not None:
    fonts = get_fonts_from_docx(uploaded_file)
    st.write("Fonts used in the document:")
    for font in fonts:
        st.write(font)

